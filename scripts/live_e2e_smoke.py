from __future__ import annotations

import argparse
import json
import os
import socket
import subprocess
import sys
import tempfile
import time
import urllib.error
import urllib.parse
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Sequence, TextIO

from ml_job_swarm.product_goals import (
    build_live_smoke_product_metrics,
    manual_submit_boundary_report,
)


DEFAULT_COMPANY = "Anthropic"
DEFAULT_BOARD_URL = "https://boards.greenhouse.io/anthropic"


@dataclass
class ServerHandle:
    process: subprocess.Popen[str]
    log_file: TextIO
    log_path: Path


def main(argv: Sequence[str] | None = None) -> int:
    args = _parse_args(argv)
    started_at = time.monotonic()
    artifact_dir = Path(args.artifact_dir) if args.artifact_dir else _temp_artifact_dir()
    artifact_dir.mkdir(parents=True, exist_ok=True)
    _progress(f"Artifacts: {artifact_dir}")
    seed_path = write_seed_catalog(
        artifact_dir,
        company_name=args.company,
        board_url=args.board_url,
    )
    resume_path = write_resume_docx(artifact_dir)
    port = args.port or _free_port(args.host)
    base_url = f"http://{args.host}:{port}"
    server = _start_server(
        host=args.host,
        port=port,
        db_path=artifact_dir / "browser.db",
        seed_path=seed_path,
        log_path=artifact_dir / "uvicorn.log",
    )
    try:
        _wait_for_app(base_url, timeout_seconds=args.server_timeout)
        summary = run_browser_smoke(
            base_url=base_url,
            resume_path=resume_path,
            artifact_dir=artifact_dir,
            browser_timeout_ms=args.browser_timeout_ms,
            headed=args.headed,
        )
    finally:
        _stop_server(server)

    summary["artifact_dir"] = str(artifact_dir)
    summary["seed_path"] = str(seed_path)
    summary["resume_path"] = str(resume_path)
    summary["server_log_path"] = str(server.log_path)
    manual_submit_report = manual_submit_boundary_report(
        Path(__file__).resolve().parents[1] / "ml_job_swarm"
    )
    summary["product_metrics"] = build_live_smoke_product_metrics(
        refresh_summary={
            "jobs_seen": int(summary["jobs_seen"]),
            "sources_attempted": int(summary["sources_attempted"]),
            "sources_succeeded": int(summary["sources_succeeded"]),
        },
        packet_prepared=summary["status"] == "browser_e2e_ok",
        saved_jobs_count=1,
        elapsed_seconds=round(time.monotonic() - started_at, 2),
        external_submit_paths=len(manual_submit_report["external_submit_paths"]),
    )
    print(json.dumps(summary, indent=2, sort_keys=True))
    return 0


def write_seed_catalog(
    artifact_dir: Path,
    *,
    company_name: str = DEFAULT_COMPANY,
    board_url: str = DEFAULT_BOARD_URL,
) -> Path:
    seed_path = artifact_dir / "seed.json"
    seed_path.write_text(
        json.dumps(
            [
                {
                    "name": company_name,
                    "aliases": [],
                    "tags": ["ai_lab"],
                    "stage": "growth",
                    "priority_tier": 1,
                    "careers_url": board_url,
                    "ats_type": "greenhouse",
                    "reviewed_at": "2026-05-11",
                }
            ],
            indent=2,
        ),
        encoding="utf-8",
    )
    return seed_path


def write_resume_docx(artifact_dir: Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to generate the smoke resume") from exc

    resume_path = artifact_dir / "resume.docx"
    document = Document()
    document.add_paragraph("Summary")
    document.add_paragraph(
        "Senior machine learning engineer building ranking systems and model "
        "serving platforms."
    )
    document.add_paragraph("Skills")
    document.add_paragraph(
        "Python, PyTorch, machine learning, model serving, distributed systems"
    )
    document.add_paragraph("Experience")
    document.add_paragraph(
        "Built production ML systems for search, recommendations, and AI "
        "infrastructure."
    )
    document.add_paragraph("Education")
    document.add_paragraph("BS Computer Science")
    document.save(resume_path)
    return resume_path


def run_browser_smoke(
    *,
    base_url: str,
    resume_path: Path,
    artifact_dir: Path,
    browser_timeout_ms: int,
    headed: bool = False,
) -> dict[str, object]:
    try:
        from playwright.sync_api import expect, sync_playwright
    except ImportError as exc:
        raise RuntimeError(
            "Playwright is required. Run with: "
            "uv run --with uvicorn --with playwright python scripts/live_e2e_smoke.py"
        ) from exc

    with sync_playwright() as playwright:
        chrome = _local_chrome_executable()
        browser = playwright.chromium.launch(
            executable_path=str(chrome) if chrome else None,
            headless=not headed,
        )
        page = browser.new_page(viewport={"width": 1440, "height": 1000})
        page.set_default_timeout(browser_timeout_ms)
        page.set_default_navigation_timeout(browser_timeout_ms)
        _progress("Opening onboarding")
        page.goto(f"{base_url}/onboarding", wait_until="domcontentloaded")
        _progress("Uploading resume")
        page.set_input_files('input[name="resume"]', str(resume_path))
        page.click('form[action="/resume"] button[type="submit"]')
        expect(page).to_have_url(_path_pattern("/onboarding?resume_asset_id="))

        _progress("Saving target profile")
        page.fill('input[name="role"]', "Machine Learning Engineer")
        page.fill('input[name="level"]', "Senior")
        page.fill('input[name="location"]', "Remote US")
        page.fill('input[name="work_mode"]', "remote")
        page.fill('input[name="company_stage"]', "growth")
        page.click('form[action="/preferences"] button[type="submit"]')
        expect(page).to_have_url(_path_pattern("/dashboard?target_profile_id="))

        _progress("Refreshing public source")
        page.click('form[action="/dashboard/refresh-sources"] button[type="submit"]')
        expect(page).to_have_url(_contains_pattern("refresh_status=completed"))
        refresh_query = urllib.parse.urlsplit(page.url).query
        refresh_summary = parse_refresh_summary(refresh_query)
        if int(refresh_summary["jobs_seen"]) <= 0:
            raise AssertionError(f"Expected live jobs_seen > 0, got {refresh_summary}")
        expect(page.locator("body")).to_contain_text("Sources attempted")
        expect(page.locator("body")).to_contain_text("Sources succeeded")
        expect(page.locator("body")).to_contain_text("Rules preview")
        page.screenshot(path=str(artifact_dir / "dashboard-after-refresh.png"), full_page=True)

        _progress("Saving first matching job")
        page.locator('form:has(input[name="decision"][value="saved"])').first.locator(
            "button"
        ).click()
        expect(page).to_have_url(_contains_pattern("decision_status=saved"))
        _progress("Opening saved jobs")
        page.click('a[href^="/dashboard/saved?target_profile_id="]')
        expect(page).to_have_url(_path_pattern("/dashboard/saved?target_profile_id="))
        expect(page.locator("body")).to_contain_text("Not reviewed")
        expect(page.locator("body")).to_contain_text("Prepare application")
        page.screenshot(path=str(artifact_dir / "saved-jobs.png"), full_page=True)

        _progress("Preparing application packet")
        page.click('form[action^="/jobs/"][action$="/application-packet"] button[type="submit"]')
        expect(page).to_have_url(_path_pattern("/jobs/"))
        expect(page.locator("body")).to_contain_text("Status: prepared")
        expect(page.locator("body")).to_contain_text("Manual submit")
        page.screenshot(path=str(artifact_dir / "job-prepared.png"), full_page=True)
        browser.close()

    return {
        "status": "browser_e2e_ok",
        "base_url": base_url,
        "jobs_seen": refresh_summary["jobs_seen"],
        "sources_attempted": refresh_summary["sources_attempted"],
        "sources_succeeded": refresh_summary["sources_succeeded"],
        "dashboard_screenshot": str(artifact_dir / "dashboard-after-refresh.png"),
        "saved_jobs_screenshot": str(artifact_dir / "saved-jobs.png"),
        "prepared_screenshot": str(artifact_dir / "job-prepared.png"),
    }


def parse_refresh_summary(query: str) -> dict[str, int]:
    params = urllib.parse.parse_qs(query)
    keys = ("jobs_seen", "sources_attempted", "sources_succeeded")
    return {key: int(params.get(key, ["0"])[0] or 0) for key in keys}


def _start_server(
    *,
    host: str,
    port: int,
    db_path: Path,
    seed_path: Path,
    log_path: Path,
) -> ServerHandle:
    env = os.environ.copy()
    env["ML_JOB_SWARM_DB_PATH"] = str(db_path)
    env["ML_JOB_SWARM_SEED_COMPANIES"] = str(seed_path)
    log_file = log_path.open("w", encoding="utf-8")
    try:
        process: subprocess.Popen[str] = subprocess.Popen(
            [
                sys.executable,
                "-m",
                "uvicorn",
                "ml_job_swarm.app:create_app_from_env",
                "--factory",
                "--host",
                host,
                "--port",
                str(port),
            ],
            env=env,
            stdout=log_file,
            stderr=subprocess.STDOUT,
            text=True,
        )
    except Exception:
        log_file.close()
        raise
    return ServerHandle(process=process, log_file=log_file, log_path=log_path)


def _wait_for_app(base_url: str, *, timeout_seconds: float) -> None:
    deadline = time.monotonic() + timeout_seconds
    last_error: Exception | None = None
    while time.monotonic() < deadline:
        try:
            with urllib.request.urlopen(f"{base_url}/onboarding", timeout=2) as response:
                if response.status == 200:
                    return
        except (urllib.error.URLError, TimeoutError) as exc:
            last_error = exc
        time.sleep(0.25)
    raise RuntimeError(f"App did not start at {base_url}: {last_error}")


def _stop_server(server: ServerHandle) -> None:
    try:
        if server.process.poll() is not None:
            return
        server.process.terminate()
        try:
            server.process.wait(timeout=5)
        except subprocess.TimeoutExpired:
            server.process.kill()
            server.process.wait(timeout=5)
    finally:
        server.log_file.close()


def _free_port(host: str) -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind((host, 0))
        return int(sock.getsockname()[1])


def _temp_artifact_dir() -> Path:
    return Path(tempfile.mkdtemp(prefix="ml-job-swarm-live-e2e."))


def _progress(message: str) -> None:
    print(f"[live-e2e] {message}", file=sys.stderr, flush=True)


def _local_chrome_executable() -> Path | None:
    chrome = Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome")
    return chrome if chrome.exists() else None


def _path_pattern(path_start: str):
    import re

    return re.compile(re.escape(path_start))


def _contains_pattern(value: str):
    import re

    return re.compile(re.escape(value))


def _parse_args(argv: Sequence[str] | None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Run a live no-credentials browser smoke against public ATS data."
    )
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=0, help="0 chooses a free port")
    parser.add_argument("--company", default=DEFAULT_COMPANY)
    parser.add_argument("--board-url", default=DEFAULT_BOARD_URL)
    parser.add_argument("--artifact-dir")
    parser.add_argument("--headed", action="store_true")
    parser.add_argument("--server-timeout", type=float, default=30)
    parser.add_argument("--browser-timeout-ms", type=int, default=60_000)
    return parser.parse_args(argv)


if __name__ == "__main__":
    raise SystemExit(main())
